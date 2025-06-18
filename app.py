import streamlit as st
import PyPDF2
from docx import Document
import io
import numpy as np
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import pandas as pd
import requests
import tempfile
import os
import json

# 페이지 설정
st.set_page_config(
    page_title="AHN'S AI Assistant",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 깔끔한 화이트 모드 CSS
st.markdown("""
<style>
    /* 메인 컨테이너 */
    .main .block-container {
        padding-top: 2rem;
        padding-bottom: 120px;
        padding-left: 2rem;
        padding-right: 2rem;
        max-width: none;
    }
    
    /* 사이드바 스타일 */
    .css-1d391kg {
        background-color: #f8f9fa;
        border-right: 1px solid #e9ecef;
    }
    
    /* 헤더 스타일 */
    h1 {
        color: #2c3e50;
        font-weight: 700;
        margin-bottom: 0.5rem;
    }
    
    h2, h3 {
        color: #34495e;
        font-weight: 600;
    }
    
    /* 버튼 스타일 */
    .stButton > button {
        background-color: #3498db;
        color: white;
        border: none;
        border-radius: 6px;
        font-weight: 500;
        padding: 0.6rem 1.2rem;
        transition: all 0.2s ease;
        box-shadow: 0 2px 4px rgba(52, 152, 219, 0.2);
    }
    
    .stButton > button:hover {
        background-color: #2980b9;
        transform: translateY(-1px);
        box-shadow: 0 4px 8px rgba(52, 152, 219, 0.3);
    }
    
    /* Primary 버튼 */
    .stButton > button[kind="primary"] {
        background-color: #27ae60;
        box-shadow: 0 2px 4px rgba(39, 174, 96, 0.2);
    }
    
    .stButton > button[kind="primary"]:hover {
        background-color: #229954;
        box-shadow: 0 4px 8px rgba(39, 174, 96, 0.3);
    }
    
    /* 입력 필드 스타일 */
    .stTextInput > div > div > input {
        border: 2px solid #e9ecef;
        border-radius: 6px;
        padding: 0.6rem;
        transition: border-color 0.2s ease;
    }
    
    .stTextInput > div > div > input:focus {
        border-color: #3498db;
        box-shadow: 0 0 0 3px rgba(52, 152, 219, 0.1);
    }
    
    /* 파일 업로더 스타일 */
    .stFileUploader {
        border: 2px dashed #bdc3c7;
        border-radius: 8px;
        padding: 1.5rem;
        background-color: #f8f9fa;
        transition: all 0.2s ease;
    }
    
    .stFileUploader:hover {
        border-color: #3498db;
        background-color: #ecf0f1;
    }
    
    /* 메시지 스타일 */
    .stSuccess {
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        color: #155724;
        border-radius: 6px;
    }
    
    .stInfo {
        background-color: #d1ecf1;
        border: 1px solid #bee5eb;
        color: #0c5460;
        border-radius: 6px;
    }
    
    .stWarning {
        background-color: #fff3cd;
        border: 1px solid #ffeaa7;
        color: #856404;
        border-radius: 6px;
    }
    
    .stError {
        background-color: #f8d7da;
        border: 1px solid #f5c6cb;
        color: #721c24;
        border-radius: 6px;
    }
    
    /* 확장 가능한 섹션 */
    .streamlit-expanderHeader {
        background-color: #f8f9fa;
        border: 1px solid #e9ecef;
        border-radius: 6px;
        font-weight: 500;
    }
    
    .streamlit-expanderContent {
        border: 1px solid #e9ecef;
        border-top: none;
        border-radius: 0 0 6px 6px;
        background-color: #ffffff;
    }
    
    /* 메트릭 컨테이너 */
    [data-testid="metric-container"] {
        background-color: #ffffff;
        border: 1px solid #e9ecef;
        border-radius: 6px;
        padding: 1rem;
        box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
    }
    
    /* 구분선 */
    hr {
        border-color: #e9ecef;
        margin: 1.5rem 0;
    }
    
    /* 채팅 입력창 스타일 수정 */
    [data-testid="stChatInput"] textarea {
        border: 2px solid #e9ecef !important;
        border-radius: 25px !important;
        padding: 12px 20px !important;
        font-size: 1rem !important;
        transition: border-color 0.2s ease !important;
        outline: none !important;
        box-shadow: none !important;
    }
    
    [data-testid="stChatInput"] textarea:focus {
        border-color: #3498db !important;
        box-shadow: 0 0 0 3px rgba(52, 152, 219, 0.1) !important;
        outline: none !important;
    }
    
    [data-testid="stChatInput"] textarea::placeholder {
        color: #7f8c8d !important;
    }
    
    /* 빨간 밑줄 및 테두리 제거 */
    [data-testid="stChatInput"] textarea:invalid {
        border-color: #e9ecef !important;
        box-shadow: none !important;
    }
    
    /* 포커스 시 빨간 테두리 방지 */
    [data-testid="stChatInput"] textarea:focus:invalid {
        border-color: #3498db !important;
        box-shadow: 0 0 0 3px rgba(52, 152, 219, 0.1) !important;
    }
    
    /* 채팅 입력창 위치 */
    .stChatInput {
        position: fixed;
        bottom: 0;
        left: 320px;
        right: 0;
        background: white;
        border-top: 1px solid #e9ecef;
        padding: 1rem 2rem;
        z-index: 999;
        box-shadow: 0 -2px 10px rgba(0,0,0,0.1);
    }
    
    /* 사이드바 축소시 */
    @media (max-width: 768px) {
        .stChatInput {
            left: 0;
            padding: 1rem;
        }
        
        .main .block-container {
            padding-left: 1rem;
            padding-right: 1rem;
        }
    }
    
    /* 링크 색상 */
    a {
        color: #3498db;
        text-decoration: none;
    }
    
    a:hover {
        color: #2980b9;
        text-decoration: underline;
    }
</style>
""", unsafe_allow_html=True)

# 제목 및 헤더
st.title("AHN'S AI Assistant")
st.markdown("**Enterprise Document Intelligence Platform**")
st.markdown("---")

def get_fallback_response(prompt: str, context: str = "", error_msg: str = "") -> str:
    """API 실패시 폴백 응답"""
    
    # 자주 묻는 질문에 대한 직접 답변
    prompt_lower = prompt.lower()
    
    if any(word in prompt_lower for word in ['와이파이', 'wifi', '무선']):
        if context and 'pstorm' in context.lower():
            return """📶 **와이파이 정보**

네트워크명: Pstorm_Office
비밀번호: Pstorm#2023
ID: pstorm2019@gmail.com

보안: WPA2-PSK
대역폭: 2.4GHz/5GHz 듀얼밴드

💡 위 정보로 무선 네트워크에 연결하세요!"""
    
    elif any(word in prompt_lower for word in ['adobe', '어도비']):
        return """🎨 **Adobe 관련 정보**

Adobe 계정이나 라이선스 정보를 찾고 계시는군요.
문서에서 관련 정보를 확인해 보겠습니다.

💡 더 정확한 정보를 위해 문서를 업데이트하거나 관리자에게 문의하세요."""
    
    elif any(word in prompt_lower for word in ['gmail', '구글', 'google']):
        return """📧 **Gmail 관련 정보**

Gmail 계정 정보를 찾고 계시는군요.
보안을 위해 정확한 계정 정보는 문서를 직접 확인해주세요.

💡 문서에서 관련 정보를 검색하거나 관리자에게 문의하세요."""
    
    # 컨텍스트가 있으면 간단한 답변 제공
    elif context.strip():
        return f"""📋 **문서 기반 답변**

질문: {prompt}

관련 문서 내용:
{context[:300]}{'...' if len(context) > 300 else ''}

⚠️ AI 서비스 연결에 문제가 있어 간단한 정보만 제공합니다.
더 자세한 분석을 원하시면 잠시 후 다시 시도해주세요.

오류 정보: {error_msg}"""
    
    # 일반적인 응답
    else:
        return f"""🤖 **시스템 알림**

죄송합니다. 현재 AI 응답 서비스에 일시적인 문제가 발생했습니다.

**문제**: {error_msg}

**해결 방법**:
1. 잠시 후 다시 시도해주세요
2. 문서에서 직접 정보를 찾아보세요
3. 네트워크 연결을 확인해주세요

질문하신 내용: "{prompt}"

💡 서비스가 복구되면 더 자세한 답변을 드리겠습니다."""

# Luxia API 설정
LUXIA_API_KEY = "U2FsdGVkX19ZW0c+KOFb9zDy5eoyiz+I6icUKb2uOjuvUnzY1TaixWa5Ouy0s87vCdtqiQMmScIWcRbEJWcfXt/jS6RMWCW+38TU47bpj82JdafHt3ODi9VHfPmSrZJCMTwP4BJ471NZTqTLakFLpMQ/PTjafRebBJpfLSDeyBj4fX1VM+NnoH8u8aGG5AV4"

# Luxia API 설정 - 실제 API 키 사용
LUXIA_API_KEY = "lZTqTLakFLpMQ/PTjafRebBJpfLSDeyBj4fXiVM+NnoH8u8aGG5AV4"

def get_luxia_response(prompt: str, context: str = "") -> str:
    """Luxia API를 통한 답변 생성 - 정확한 공식 엔드포인트 사용"""
    try:
        # 공식 Luxia API 엔드포인트
        url = "https://bridge.luxiacloud.com/luxia/v1/chat"
        
        headers = {
            "apikey": LUXIA_API_KEY,  # 'Authorization: Bearer' 대신 'apikey' 사용
            "Content-Type": "application/json"
        }
        
        # 컨텍스트와 질문을 결합
        if context.strip():
            full_prompt = f"""다음 문서 정보를 참고하여 질문에 정확하고 자세하게 답변해주세요:

[문서 정보]
{context}

[질문]
{prompt}

답변은 친절하고 이해하기 쉽게 작성해주세요."""
        else:
            full_prompt = prompt

        # Luxia 공식 API 형식
        payload = {
            "model": "luxia2.5-llm-32b-0401",  # 공식 모델명
            "messages": [
                {
                    "role": "user",
                    "content": full_prompt
                }
            ]
        }
        
        st.info("🚀 Luxia API 연결 중...")
        
        response = requests.post(
            url, 
            json=payload, 
            headers=headers, 
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            
            # Luxia API 응답 처리
            if 'choices' in result and len(result['choices']) > 0:
                message_content = result['choices'][0].get('message', {}).get('content', '')
                if message_content:
                    st.success("✅ Luxia API 연결 성공!")
                    return message_content
            
            # 대체 응답 형식 처리
            elif 'message' in result:
                st.success("✅ Luxia API 연결 성공!")
                return result['message']
            
            elif 'response' in result:
                st.success("✅ Luxia API 연결 성공!")
                return result['response']
            
            else:
                st.error(f"예상치 못한 응답 형식: {result}")
                return "API 응답 형식을 인식할 수 없습니다."
        
        elif response.status_code == 401:
            st.error("🔑 API 키 인증 실패 - platform.luxiacloud.com에서 키를 확인해주세요")
            return "API 키 인증에 실패했습니다. 대시보드에서 API 키 상태를 확인해주세요."
            
        elif response.status_code == 429:
            st.warning("⏰ API 사용 한도 초과 - 잠시 후 다시 시도해주세요")
            return "API 사용 한도에 도달했습니다. 잠시 후 다시 시도해주세요."
            
        elif response.status_code == 400:
            st.error(f"❌ 요청 오류 (400): {response.text}")
            return "요청 형식에 오류가 있습니다. 관리자에게 문의하세요."
            
        else:
            st.error(f"❌ API 오류 ({response.status_code}): {response.text}")
            return f"API 오류가 발생했습니다. 상태 코드: {response.status_code}"
            
    except requests.exceptions.ConnectionError:
        st.error("🔌 네트워크 연결 오류 - 인터넷 연결을 확인해주세요")
        return "네트워크 연결에 문제가 있습니다. 인터넷 연결을 확인해주세요."
        
    except requests.exceptions.Timeout:
        st.warning("⏰ 응답 시간 초과 - 다시 시도해주세요")
        return "응답 시간이 초과되었습니다. 다시 시도해주세요."
        
    except Exception as e:
        st.error(f"❌ 시스템 오류: {str(e)}")
        return f"예상치 못한 오류가 발생했습니다: {str(e)}"

# 세션 상태 초기화
if 'messages' not in st.session_state:
    st.session_state.messages = []

if 'documents' not in st.session_state:
    st.session_state.documents = []

if 'embeddings' not in st.session_state:
    st.session_state.embeddings = None

if 'encoder' not in st.session_state:
    st.session_state.encoder = None

if 'default_loaded' not in st.session_state:
    st.session_state.default_loaded = False

# 문서 처리 함수들
def extract_text_from_pdf(file):
    """PDF에서 텍스트 추출"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"PDF reading error: {e}")
        return ""

def extract_text_from_docx(file):
    """DOCX에서 텍스트 추출"""
    try:
        doc = Document(io.BytesIO(file.read()))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"DOCX reading error: {e}")
        return ""

def extract_text_from_txt(file):
    """TXT에서 텍스트 추출"""
    try:
        return file.read().decode('utf-8')
    except Exception as e:
        st.error(f"TXT reading error: {e}")
        return ""

def split_text_into_chunks(text, chunk_size=500):
    """텍스트를 청크로 분할"""
    if not text.strip():
        return []
    
    sentences = text.replace('\n', ' ').split('. ')
    chunks = []
    current_chunk = []
    current_length = 0
    
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
            
        sentence_length = len(sentence)
        
        if current_length + sentence_length > chunk_size and current_chunk:
            chunks.append('. '.join(current_chunk) + '.')
            current_chunk = [sentence]
            current_length = sentence_length
        else:
            current_chunk.append(sentence)
            current_length += sentence_length
    
    if current_chunk:
        chunks.append('. '.join(current_chunk) + '.')
    
    return [chunk for chunk in chunks if len(chunk.strip()) > 50]

def process_documents(files):
    """업로드된 문서들 처리"""
    documents = []
    
    for file in files:
        try:
            if file.type == "application/pdf":
                text = extract_text_from_pdf(file)
            elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = extract_text_from_docx(file)
            elif file.type == "text/plain":
                text = extract_text_from_txt(file)
            else:
                st.warning(f"Unsupported file format: {file.name}")
                continue
            
            if not text.strip():
                st.warning(f"No text extracted from: {file.name}")
                continue
            
            chunks = split_text_into_chunks(text, chunk_size=500)
            
            for i, chunk in enumerate(chunks):
                documents.append({
                    'id': f"{file.name}_{i}",
                    'text': chunk,
                    'filename': file.name,
                    'chunk_id': i
                })
        
        except Exception as e:
            st.error(f"Error processing {file.name}: {e}")
    
    return documents

def fallback_keyword_search(query, documents, n_results=3):
    """임베딩 실패시 키워드 검색으로 폴백"""
    try:
        if not documents:
            return []
        
        query_lower = query.lower()
        scored_docs = []
        
        for doc in documents:
            text_lower = doc['text'].lower()
            score = 0
            
            # 키워드 매칭 점수 계산
            query_words = query_lower.split()
            for word in query_words:
                if len(word) > 1:  # 한 글자 단어 제외
                    score += text_lower.count(word)
            
            if score > 0:
                scored_docs.append((doc['text'], score))
        
        # 점수별 정렬
        scored_docs.sort(key=lambda x: x[1], reverse=True)
        
        # 상위 결과 반환
        results = [doc[0] for doc in scored_docs[:n_results]]
        
        if results:
            st.info("키워드 검색으로 대체하여 결과를 찾았습니다.")
        
        return results
    
    except Exception as e:
        st.error(f"폴백 검색 오류: {e}")
        return []

@st.cache_resource
def load_sentence_transformer():
    """SentenceTransformer 모델 로드 - 오류 처리 강화"""
    try:
        st.info("SentenceTransformer 모델을 로딩하고 있습니다...")
        model = SentenceTransformer('all-MiniLM-L6-v2')
        
        # 모델이 제대로 로드되었는지 테스트
        test_encoding = model.encode(["test"])
        if test_encoding is not None and len(test_encoding) > 0:
            st.success("✅ SentenceTransformer 모델 로드 성공!")
            return model
        else:
            raise Exception("모델 테스트 인코딩 실패")
            
    except Exception as e:
        st.error(f"SentenceTransformer 로딩 실패: {e}")
        st.warning("⚠️ 임베딩 검색이 비활성화됩니다. 키워드 검색으로 대체됩니다.")
        return None

def create_embeddings(documents):
    """문서 임베딩 생성 - 안전성 강화"""
    if not documents:
        st.warning("처리할 문서가 없습니다.")
        return None, None
    
    try:
        st.info("임베딩 모델을 로드하고 있습니다...")
        encoder = load_sentence_transformer()
        
        if encoder is None:
            st.warning("임베딩 모델 로드에 실패했습니다. 키워드 검색으로 대체됩니다.")
            return None, None
        
        st.info(f"{len(documents)}개 문서의 임베딩을 생성하고 있습니다...")
        texts = [doc['text'] for doc in documents]
        
        # 배치 단위로 처리하여 메모리 절약
        batch_size = 32
        all_embeddings = []
        
        for i in range(0, len(texts), batch_size):
            batch_texts = texts[i:i+batch_size]
            batch_embeddings = encoder.encode(batch_texts, show_progress_bar=False)
            all_embeddings.extend(batch_embeddings)
        
        embeddings = np.array(all_embeddings)
        st.success(f"✅ {len(embeddings)}개 문서 임베딩 생성 완료!")
        
        return embeddings, encoder
    
    except Exception as e:
        st.error(f"임베딩 생성 오류: {e}")
        st.warning("키워드 검색으로 대체됩니다.")
        return None, None

def search_documents(query, documents, embeddings, encoder, n_results=3):
    """문서에서 관련 내용 검색"""
    try:
        # None 체크 강화
        if not documents or embeddings is None or encoder is None:
            st.warning("검색 시스템이 초기화되지 않았습니다. 문서를 다시 처리해주세요.")
            return []
        
        # encoder가 제대로 로드되었는지 확인
        if not hasattr(encoder, 'encode'):
            st.error("SentenceTransformer 모델이 제대로 로드되지 않았습니다.")
            return []
        
        query_embedding = encoder.encode([query])
        similarities = cosine_similarity(query_embedding, embeddings)[0]
        top_indices = np.argsort(similarities)[::-1][:n_results]
        
        results = []
        for idx in top_indices:
            if similarities[idx] > 0.1:
                results.append(documents[idx]['text'])
        
        return results
    
    except Exception as e:
        st.error(f"Document search error: {e}")
        # 폴백: 키워드 검색으로 대체
        return fallback_keyword_search(query, documents, n_results)

def load_default_document():
    """GitHub에서 기본 문서 로드 (password-rag-agent2 경로 사용)"""
    try:
        # GitHub raw 파일 URL - password-rag-agent2로 변경
        github_url = "https://raw.githubusercontent.com/ppacksae/password-rag-agent2/main/pstorm_pw.docx"
        
        # 파일 다운로드
        response = requests.get(github_url, timeout=10)
        
        if response.status_code == 200:
            # 임시 파일로 저장
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(response.content)
                tmp_file_path = tmp_file.name
            
            try:
                # DOCX 파일 읽기
                doc = Document(tmp_file_path)
                content = ""
                for paragraph in doc.paragraphs:
                    content += paragraph.text + "\n"
                
                # 임시 파일 삭제
                os.unlink(tmp_file_path)
                
                if not content.strip():
                    raise Exception("문서가 비어있습니다")
                
            except Exception as e:
                # 임시 파일 삭제 (오류 발생 시에도)
                if os.path.exists(tmp_file_path):
                    os.unlink(tmp_file_path)
                raise e
        else:
            raise Exception(f"GitHub에서 파일을 가져올 수 없습니다. Status: {response.status_code}")
        
        # 텍스트를 청크로 분할
        chunks = split_text_into_chunks(content, chunk_size=500)
        
        documents = []
        for i, chunk in enumerate(chunks):
            documents.append({
                'id': f"pstorm_pw.docx_{i}",
                'text': chunk,
                'filename': "pstorm_pw.docx",
                'chunk_id': i
            })
        
        return documents
    
    except Exception as e:
        st.error(f"GitHub에서 기본 문서 로드 중 오류: {e}")
        st.info("기본 내용으로 대체합니다.")
        
        # 폴백 - 기본 내용 사용
        default_content = """
6. 와이파이(WIFI)
1) 비번(password) : Pstorm#2023
2) ID: pstorm2019@gmail.com

네트워크 설정:
- 네트워크명: Pstorm_Office
- 보안: WPA2-PSK
- 대역폭: 2.4GHz/5GHz 듀얼밴드
- 최대 연결 기기: 50대

관리자 정보:
- 관리자 ID: admin
- 관리자 비밀번호: admin123!
- 웹 관리 주소: 192.168.1.1

추가 정보:
- 게스트 네트워크: Pstorm_Guest
- 게스트 비밀번호: guest2023
- 포트 포워딩: 활성화
- 방화벽: 기본 설정
"""
        
        chunks = split_text_into_chunks(default_content, chunk_size=500)
        documents = []
        for i, chunk in enumerate(chunks):
            documents.append({
                'id': f"pstorm_pw.docx_{i}",
                'text': chunk,
                'filename': "pstorm_pw.docx",
                'chunk_id': i
            })
        
        return documents

def generate_response(query, context_docs):
    """Luxia를 사용하여 응답 생성"""
    try:
        if context_docs:
            context = "\n\n".join(context_docs)
            response = get_luxia_response(query, context)
        else:
            no_docs_prompt = f"""
업로드된 문서에서 관련 정보를 찾을 수 없습니다.

질문: {query}

일반적인 지식을 바탕으로 답변을 제공하겠습니다만, 먼저 업로드된 문서에서 관련 정보를 찾지 못했다는 점을 말씀드립니다.
"""
            response = get_luxia_response(no_docs_prompt)
        
        return response
    
    except Exception as e:
        return f"응답 생성 중 오류가 발생했습니다: {e}"

# 사이드바 설정
with st.sidebar:
    st.header("Configuration")
    
    # API 키 정보 표시 (읽기 전용)
    st.markdown("### 🔑 AI Model")
    st.success("🚀 **Luxia Platform** 연결 준비됨")
    
    st.markdown("---")
    
    # 문서 관리 섹션
    st.header("Document Management")
    
    # 문서 업로드
    st.subheader("Upload Documents")
    uploaded_files = st.file_uploader(
        "Supported formats: PDF, DOCX, TXT",
        type=['pdf', 'docx', 'txt'],
        accept_multiple_files=True,
        help="Upload company documents for AI analysis"
    )
    
    # 기본 문서 자동 로드
    if not st.session_state.default_loaded and not st.session_state.documents:
        with st.spinner("기본 문서를 로드하고 있습니다..."):
            default_docs = load_default_document()
            if default_docs:
                st.session_state.documents = default_docs
                
                # 임베딩 생성
                embeddings, encoder = create_embeddings(default_docs)
                if embeddings is not None:
                    st.session_state.embeddings = embeddings
                    st.session_state.encoder = encoder
                    st.session_state.default_loaded = True
                    st.success("✅ 기본 문서 (pstorm_pw.docx) 로드 완료!")
                    st.rerun()
    
    # 문서 처리 버튼
    if uploaded_files:
        if st.button("Process Documents", type="primary", use_container_width=True):
            with st.spinner("Processing documents..."):
                # 문서 처리 로직
                documents = process_documents(uploaded_files)
                
                if documents:
                    st.session_state.documents = documents
                    
                    with st.spinner("Generating embeddings..."):
                        embeddings, encoder = create_embeddings(documents)
                        if embeddings is not None:
                            st.session_state.embeddings = embeddings
                            st.session_state.encoder = encoder
                            st.success(f"Processed {len(documents)} document chunks")
                        else:
                            st.error("Embedding generation failed")
                else:
                    st.warning("No processable documents found")
                
                st.rerun()
    
    st.markdown("---")
    
    # 문서 현황
    st.subheader("Document Status")
    if st.session_state.get('documents'):
        st.metric("Total Chunks", len(st.session_state.documents))
        
        # 파일별 청크 수 표시
        file_counts = {}
        for doc in st.session_state.documents:
            filename = doc['filename']
            file_counts[filename] = file_counts.get(filename, 0) + 1
        
        for filename, count in file_counts.items():
            st.text(f"{filename}: {count} chunks")
        
        # 검색 기능 상태
        if st.session_state.get('embeddings') is not None:
            st.success("🔍 Vector Search: Active")
        elif st.session_state.get('documents'):
            st.info("🔤 Keyword Search: Active")
        else:
            st.warning("❌ Search: Inactive")
    else:
        st.info("No documents loaded")
    
    st.markdown("---")
    
    # 관리 기능
    st.subheader("System Management")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("Clear Chat", use_container_width=True):
            st.session_state.messages = []
            st.rerun()
    
    with col2:
        if st.button("Clear Docs", use_container_width=True):
            st.session_state.documents = []
            st.session_state.embeddings = None
            st.session_state.encoder = None
            st.rerun()

# 사용법 안내를 사이드바로 이동
with st.sidebar:
    st.markdown("---")
    st.subheader("System Information")
    
    with st.expander("Getting Started"):
        st.markdown("""
        1. AI is powered by Luxia (already connected)
        2. Upload documents using the file uploader
        3. Click "Process Documents" to enable AI search
        4. Ask questions about your documents in the chat
        """)
    
    with st.expander("Features"):
        st.markdown("""
        - PDF, DOCX, TXT file support
        - Multiple file upload capability
        - Vector-based document search
        - Professional AI responses
        - Source document references
        """)
    
    with st.expander("Tips"):
        st.markdown("""
        - Use specific questions for better results
        - Multiple files can be processed together
        - First document processing may take time
        - Referenced sources shown below responses
        """)

# 메인 채팅 인터페이스
st.header("AI Chat Interface")

# 채팅 컨테이너 (사이드바 너비만큼 여백 추가)
chat_container = st.container()

with chat_container:
    # 초기 환영 메시지 (채팅이 비어있을 때만 표시)
    if not st.session_state.messages:
        st.markdown("""
        <div style="
            display: flex;
            justify-content: center;
            align-items: center;
            height: 300px;
            flex-direction: column;
            margin-left: 0;
        ">
            <h1 style="
                font-size: 3rem;
                font-weight: 300;
                color: #2c3e50;
                margin-bottom: 2rem;
                text-align: center;
            ">안녕하세요</h1>
            <p style="
                font-size: 1.2rem;
                color: #7f8c8d;
                text-align: center;
                margin-bottom: 3rem;
            ">AHN'S AI Assistant가 도와드리겠습니다</p>
        </div>
        """, unsafe_allow_html=True)

    # 채팅 메시지 표시 (커스텀 스타일)
    if st.session_state.messages:
        for i, message in enumerate(st.session_state.messages):
            if message["role"] == "user":
                # 사용자 메시지 - 우측 정렬
                st.markdown(f"""
                <div style="
                    display: flex;
                    justify-content: flex-end;
                    margin: 1rem 0;
                    padding-right: 1rem;
                ">
                    <div style="
                        background-color: #e3f2fd;
                        color: #1565c0;
                        padding: 0.8rem 1.2rem;
                        border-radius: 18px 18px 4px 18px;
                        max-width: 70%;
                        font-size: 0.95rem;
                        line-height: 1.4;
                        box-shadow: 0 1px 2px rgba(0,0,0,0.1);
                        word-wrap: break-word;
                    ">
                        {message["content"]}
                    </div>
                </div>
                """, unsafe_allow_html=True)
            else:
                # AI 응답 - 좌측 정렬
                # HTML 태그 제거하고 깔끔하게 표시
                clean_content = message["content"].replace('<div>', '').replace('</div>', '').strip()
                
                st.markdown(f"""
                <div style="
                    display: flex;
                    justify-content: flex-start;
                    margin: 1rem 0;
                    align-items: flex-start;
                    padding-left: 1rem;
                ">
                    <div style="
                        background-color: #f5f5f5;
                        color: #2c3e50;
                        padding: 0.8rem 1.2rem;
                        border-radius: 18px 18px 18px 4px;
                        max-width: 75%;
                        font-size: 0.95rem;
                        line-height: 1.5;
                        box-shadow: 0 1px 2px rgba(0,0,0,0.1);
                        border: 1px solid #e9ecef;
                        word-wrap: break-word;
                        white-space: pre-wrap;
                    ">
                        🤖 {clean_content}
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # 참고 문서가 있을 경우 표시
                if "references" in message:
                    with st.expander(f"📚 참고한 문서 ({len(message['references'])}개)"):
                        for j, doc in enumerate(message["references"]):
                            st.write(f"**문서 {j+1}:**")
                            st.write(doc[:200] + "..." if len(doc) > 200 else doc)
                            if j < len(message["references"]) - 1:
                                st.markdown("---")

# 커스텀 채팅 입력창
st.markdown("""
<style>
    /* 메인 컨텐츠 영역 사이드바 겹침 방지 */
    .main .block-container {
        padding-bottom: 120px !important;
        padding-left: 1rem !important;
        padding-right: 1rem !important;
    }
    
    /* 사이드바가 있을 때 메인 콘텐츠 여백 조정 */
    .main {
        margin-left: 0 !important;
    }
    
    /* 채팅 컨테이너 스타일 */
    .chat-container {
        margin-left: 0;
        width: 100%;
        max-width: none;
    }
    
    /* 채팅 입력창 커스터마이징 */
    .stChatInput {
        position: fixed;
        bottom: 0;
        left: 320px; /* 사이드바 너비만큼 여백 */
        right: 0;
        background: white;
        border-top: 1px solid #e9ecef;
        padding: 1rem;
        z-index: 999;
        box-shadow: 0 -2px 10px rgba(0,0,0,0.1);
    }
    
    /* 사이드바가 축소된 경우 */
    .css-1lcbmhc.e1fqkh3o0 + .main .stChatInput {
        left: 60px;
    }
    
    [data-testid="stChatInput"] {
        margin-bottom: 0;
        max-width: calc(100vw - 360px); /* 사이드바 고려한 최대 너비 */
    }
    
    [data-testid="stChatInput"] textarea {
        border: 2px solid #e9ecef !important;
        border-radius: 25px !important;
        padding: 12px 20px !important;
        font-size: 1rem !important;
        resize: none !important;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1) !important;
        transition: all 0.2s ease !important;
        width: 100% !important;
    }
    
    [data-testid="stChatInput"] textarea:focus {
        border-color: #3498db !important;
        box-shadow: 0 0 0 3px rgba(52, 152, 219, 0.1), 0 2px 10px rgba(0,0,0,0.15) !important;
        outline: none !important;
    }
    
    [data-testid="stChatInput"] textarea::placeholder {
        color: #7f8c8d !important;
        font-size: 1rem !important;
    }
    
    /* 채팅 메시지 영역 여백 */
    .element-container:has([data-testid="stChatInput"]) {
        margin-bottom: 80px;
    }
    
    /* 모바일 대응 */
    @media (max-width: 768px) {
        .stChatInput {
            left: 0;
        }
        [data-testid="stChatInput"] {
            max-width: 100vw;
        }
    }
</style>
""", unsafe_allow_html=True)

# 사용자 입력 (커스텀 placeholder)
if prompt := st.chat_input("AHN'S AI 에게 물어보기"):
    # 사용자 메시지 추가
    st.session_state.messages.append({"role": "user", "content": prompt})
    
    # AI 응답 생성
    with st.spinner("Luxia AI가 답변을 생성하고 있습니다..."):
        # 문서 검색
        relevant_docs = search_documents(
            prompt, 
            st.session_state.documents, 
            st.session_state.embeddings, 
            st.session_state.encoder
        )
        
        # 응답 생성
        response = generate_response(prompt, relevant_docs)
        
        # 응답을 세션에 저장 (참고 문서 포함)
        message_data = {"role": "assistant", "content": response}
        if relevant_docs:
            message_data["references"] = relevant_docs
        
        st.session_state.messages.append(message_data)
        
    # 페이지 새로고침으로 UI 업데이트
    st.rerun()

# 푸터
st.markdown("---")
st.markdown("**AHN'S AI Assistant** | Enterprise Document Intelligence Platform | Powered by Luxia AI")